#!/usr/bin/env python3
"""Convert SafeLand markdown documents to .docx files."""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUTPUT_DIR = os.path.expanduser("~/docs-safeland")
DOCS_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "docs")
B2G_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "B2G")

# Files to convert
FILES = {
    "docs": [
        ("CDC-COMPLET-V3.md", "CDC-Complet-SafeLand-V3"),
        ("LIVRE-BLANC-V2.md", "Livre-Blanc-SafeLand-V2"),
        ("LIVRE-MARKETING.md", "Livre-Marketing-SafeLand"),
        ("GUIDE-UTILISATEURS.md", "Guide-Utilisateurs-SafeLand"),
        ("AUDIT-TECHNIQUE-COMPLET.md", "Audit-Technique-SafeLand"),
    ],
    "B2G": [
        ("00-INDEX-DOSSIER-B2G.md", "B2G-00-Index-Dossier"),
        ("01-NOTE-CADRAGE-ANCFCC.md", "B2G-01-Note-Cadrage-ANCFCC"),
        ("02-DOSSIER-ADD.md", "B2G-02-Dossier-ADD"),
        ("03-CONVENTION-PPP.md", "B2G-03-Convention-PPP"),
        ("04-DECLARATION-CNDP.md", "B2G-04-Declaration-CNDP"),
        ("05-QUALIFICATION-MTNRA.md", "B2G-05-Qualification-MTNRA"),
        ("06-NOTE-DGSSI.md", "B2G-06-Note-DGSSI"),
        ("07-PROPOSITION-PILOTE.md", "B2G-07-Proposition-Pilote"),
        ("08-MARCHES-PUBLICS.md", "B2G-08-Marches-Publics"),
        ("09-MODELE-ECONOMIQUE.md", "B2G-09-Modele-Economique"),
    ],
}


def md_to_docx(md_path, output_name):
    """Convert a markdown file to a .docx file."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Read markdown
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    in_table = False
    table_rows = []
    in_code = False
    code_lines = []

    while i < len(lines):
        line = lines[i].rstrip("\n")

        # Code blocks
        if line.startswith("```"):
            if in_code:
                # End code block
                code_text = "\n".join(code_lines)
                p = doc.add_paragraph()
                p.style = doc.styles["Normal"]
                run = p.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Table rows
        if "|" in line and line.strip().startswith("|"):
            stripped = line.strip()
            # Skip separator rows like |---|---|
            if re.match(r"^\|[\s\-:|]+\|$", stripped):
                i += 1
                continue
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            # End of table
            if table_rows:
                try:
                    num_cols = max(len(r) for r in table_rows)
                    table = doc.add_table(rows=len(table_rows), cols=num_cols)
                    table.style = "Light Grid Accent 1"
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    for ri, row in enumerate(table_rows):
                        for ci, cell in enumerate(row):
                            if ci < num_cols:
                                table.rows[ri].cells[ci].text = cell.replace("**", "").replace("*", "")
                except Exception:
                    # Fallback: just add as text
                    for row in table_rows:
                        doc.add_paragraph(" | ".join(row))
            table_rows = []
            in_table = False
            # Don't increment i, process current line

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Headings
        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=0)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
            i += 1
            continue

        # Horizontal rule
        if line.strip() == "---":
            doc.add_paragraph("_" * 60)
            i += 1
            continue

        # Bullet points
        if line.startswith("- ") or line.startswith("* "):
            text = line[2:].strip()
            text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)  # Remove bold markers
            text = re.sub(r"\*(.*?)\*", r"\1", text)  # Remove italic markers
            doc.add_paragraph(text, style="List Bullet")
            i += 1
            continue

        # Numbered lists
        m = re.match(r"^(\d+)\.\s+(.+)", line)
        if m:
            text = m.group(2).strip()
            text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
            doc.add_paragraph(text, style="List Number")
            i += 1
            continue

        # Blockquote
        if line.startswith("> "):
            text = line[2:].strip()
            text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
            p = doc.add_paragraph(text)
            p.paragraph_format.left_indent = Inches(0.5)
            p.runs[0].italic = True if p.runs else None
            i += 1
            continue

        # Checkbox
        if line.startswith("- [ ] ") or line.startswith("- [x] "):
            checked = line.startswith("- [x] ")
            text = line[6:].strip()
            prefix = "[X] " if checked else "[ ] "
            doc.add_paragraph(prefix + text, style="List Bullet")
            i += 1
            continue

        # Regular paragraph
        text = line.strip()
        text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
        text = re.sub(r"\*(.*?)\*", r"\1", text)
        if text:
            doc.add_paragraph(text)

        i += 1

    # Handle any remaining table
    if in_table and table_rows:
        try:
            num_cols = max(len(r) for r in table_rows)
            table = doc.add_table(rows=len(table_rows), cols=num_cols)
            table.style = "Light Grid Accent 1"
            for ri, row in enumerate(table_rows):
                for ci, cell in enumerate(row):
                    if ci < num_cols:
                        table.rows[ri].cells[ci].text = cell.replace("**", "")
        except Exception:
            pass

    # Save .docx
    docx_path = os.path.join(OUTPUT_DIR, f"{output_name}.docx")
    doc.save(docx_path)
    print(f"  [OK] {docx_path}")
    return docx_path


def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    print(f"Output directory: {OUTPUT_DIR}")
    print()

    # Convert docs/
    print("=== Documents principaux (docs/) ===")
    for filename, output_name in FILES["docs"]:
        md_path = os.path.join(DOCS_DIR, filename)
        if os.path.exists(md_path):
            md_to_docx(md_path, output_name)
        else:
            print(f"  [SKIP] {md_path} not found")

    print()

    # Convert B2G/
    print("=== Dossier B2G ===")
    for filename, output_name in FILES["B2G"]:
        md_path = os.path.join(B2G_DIR, filename)
        if os.path.exists(md_path):
            md_to_docx(md_path, output_name)
        else:
            print(f"  [SKIP] {md_path} not found")

    print()
    print(f"Done! {len(FILES['docs']) + len(FILES['B2G'])} documents converted to {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
